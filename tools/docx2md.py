"""tools/docx2md.py -- TM Forum use-case DOCX -> knowledge/use-cases/<ID>/<ID>.md

Reaches into the <w:sdt> content control that wraps the entire body in TM
Forum's use-case template (python-docx's flat .paragraphs/.tables API sees
almost nothing in these documents without this -- see spec/spec.md 5.1),
strips the fixed set of boilerplate sections documented in spec/spec.md
5.1.1 (title block, Notice, Table of Contents, References, Administrative
Appendix and its children), and emits the universal envelope (spec/spec.md
5.0) as YAML frontmatter for every field derivable from the DOCX alone.

Fields that need the document's catalog page rather than the DOCX itself
(status, maturity, approval_status, release_status, dates, source.origin,
source.license) are left as explicit TODO stubs -- tools/add_usecase_metadata.py
fills them in afterwards without disturbing anything this script wrote
(spec/tasks.md 0.6).

Usage:
    python docx2md.py <ID> <src.docx> <out.md> <media_dir_name>

Example:
    python docx2md.py TMFS001 references/use-cases/TMFS001/TMFS001_v5.0.5.docx \
        knowledge/use-cases/TMFS001/TMFS001.md media
"""
import datetime as _dt
import hashlib
import os
import re
import sys

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

import _yaml_lite as yaml_lite

ID = sys.argv[1]
SRC = sys.argv[2]
OUT = sys.argv[3]
MEDIA_DIR_NAME = sys.argv[4]
MEDIA_DIR = os.path.join(os.path.dirname(OUT), MEDIA_DIR_NAME)

os.makedirs(MEDIA_DIR, exist_ok=True)

document = docx.Document(SRC)
body = document.element.body

# --- boilerplate to strip, spec/spec.md 5.1.1 -- fixed heading-name list, not a heuristic ---
DROP_HEADINGS = {
    "notice",
    "table of contents",
    "references",
    "administrative appendix",
    "document history",
    "version history",
    "release history",
    "acknowledgments",
    "acknowledgements",  # TM Forum spells it both ways across documents
}
HEADING_LEVEL = {"Heading1": 1, "Heading2": 2, "Heading3": 3, "Heading4": 4}
SKIP_STYLES = {"TOC1", "TOC2", "TOC3", "TOCHeading"}  # generated TOC entries -- always empty, never useful

img_counter = 0
img_map = {}  # rId -> saved relative path


def save_image(rId):
    global img_counter
    if rId in img_map:
        return img_map[rId]
    part = document.part.related_parts[rId]
    ext = part.partname.ext
    img_counter += 1
    fname = f"image{img_counter:02d}.{ext}"
    path = os.path.join(MEDIA_DIR, fname)
    with open(path, "wb") as f:
        f.write(part.blob)
    rel = f"{MEDIA_DIR_NAME}/{fname}"
    img_map[rId] = rel
    return rel


def strip_md(s):
    """Plain text for matching/regexing -- drop the bold/italic markers run_markdown() adds."""
    return re.sub(r"\*+", "", s or "").strip()


def run_markdown(run):
    """Render a run's text plus any inline images, with bold/italic markers."""
    out = []
    t = run.text
    if t:
        txt = t
        if run.bold and run.italic:
            txt = f"***{txt}***"
        elif run.bold:
            txt = f"**{txt}**"
        elif run.italic:
            txt = f"*{txt}*"
        out.append(txt)
    for blip in run._element.findall(".//" + qn("a:blip")):
        rId = blip.get(qn("r:embed"))
        if rId:
            out.append(f"\n\n![]({save_image(rId)})\n")
    return "".join(out)


def paragraph_text_and_style(p: Paragraph):
    """Return (markdown_text, style_id, is_list) for one paragraph, hyperlinks included.

    python-docx's p.runs doesn't include hyperlink runs (older versions), so we walk
    the raw XML children of the paragraph in document order instead of using p.runs.
    """
    style_id_el = p._p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    style_id = style_id_el.get(qn("w:val")) if style_id_el is not None else None

    pieces = []
    for child in p._p:
        tag = child.tag.split("}")[-1]
        if tag == "r":
            pieces.append(run_markdown(docx.text.run.Run(child, p)))
        elif tag == "hyperlink":
            rId = child.get(qn("r:id"))
            url = document.part.rels[rId].target_ref if rId and rId in document.part.rels else None
            text = "".join(docx.text.run.Run(rc, p).text or "" for rc in child.findall(qn("w:r")))
            if url and text.strip():
                pieces.append(f"[{text}]({url})")
            elif text.strip():
                pieces.append(text)
    text = "".join(pieces).strip("\n")

    numPr = p._p.find(qn("w:pPr") + "/" + qn("w:numPr"))
    is_list = numPr is not None
    return text, style_id, is_list


def paragraph_markdown(text, style_id, is_list):
    if style_id in ("Heading1", "Heading2", "Heading3", "Heading4"):
        return f"{'#' * HEADING_LEVEL[style_id]} {text}"
    if style_id == "Titlesubtitle":
        return f"### {text}"
    if is_list:
        return f"- {text}"
    return text


def table_markdown(tbl: Table):
    rows_out = []
    cells_flat = []  # (row_idx, cell_text) for callers that need to scan cell content, e.g. "Version 5.0.5"
    for i, row in enumerate(tbl.rows):
        cells = []
        for cell in row.cells:
            parts = []
            for pp in cell.paragraphs:
                txt = (pp.text or "").strip()
                if txt:
                    parts.append(txt)
                for blip in pp._p.findall(".//" + qn("a:blip")):
                    rId = blip.get(qn("r:embed"))
                    if rId:
                        parts.append(f"<br>![]({save_image(rId)})")
            cell_text = " ".join(parts).replace("|", "\\|").replace("\n", " ")
            cells.append(cell_text if cell_text else " ")
            cells_flat.append(cell_text)
        rows_out.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows_out.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows_out), cells_flat


def flatten(container_el):
    """Yield ('p', Paragraph) / ('tbl', Table) in document order, recursing into any
    w:sdt wrapper it meets at any depth -- this is what reaches past the content
    control that hides the real body from python-docx's flat API (spec/spec.md 5.1)."""
    for child in container_el:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield "p", Paragraph(child, document)
        elif tag == "tbl":
            yield "tbl", Table(child, document)
        elif tag == "sdt":
            sdt_content = child.find(qn("w:sdtContent"))
            if sdt_content is not None:
                yield from flatten(sdt_content)
        # bookmarkStart/End and others: ignore


# --- envelope fields this script can derive from the DOCX alone ---
envelope = {
    "id": ID,
    "type": "use-case",
    "name": None,  # from the "Use Case: <name>" title-block line
    "version": None,  # from the title-block table's "Version x.y.z" cell
    "status": "TODO",  # needs the catalog page -- tools/add_usecase_metadata.py
    "source": {
        "origin": "TODO",  # catalog URL -- needs the catalog page, see add_usecase_metadata.py
        "license": "TODO",  # ditto
        "retrieved": _dt.date.fromtimestamp(os.path.getmtime(SRC)).isoformat(),  # when SRC was saved to disk
        "sha256": hashlib.sha256(open(SRC, "rb").read()).hexdigest(),
        "raw_path": SRC,
    },
    "links": {"components": [], "apis": [], "use_cases": []},
}
extension = {
    "maturity": "TODO",
    "approval_status": "TODO",
    "release_status": "TODO",
    "team_approved": "TODO",
    "published": "TODO",
    "sid_references": [],
}

# --- walk the body once: title block -> strip+parse; boilerplate sections -> strip; everything else -> keep ---
lines = []
seen_first_heading = False
drop_until_level = None  # None = keep everything; int = we're inside a dropped subtree at this level
in_references = False
ref_block = None  # "sid" | "components" | "apis" | None, while in_references
title_candidate = None  # best guess at the name line while scanning the title block, see below

# The References section is NOT uniformly formatted across TM Forum use cases -- confirmed
# by checking all six pilot documents (spec/spec.md 5.1.1 assumed it was; it isn't). The three
# GA/TM-Forum-Approved documents (TMFS001/002/009) use a clean "TMFCxxx Name vX.Y.Z" bullet
# format; the three Beta/Alpha ones (TMFS029/030/031) each phrase it differently -- ID-first-
# no-version, "Name - TMFCxxx", "Name (TMFxxx)". This tracks the same maturity finding from
# the research phase: less mature documents haven't converged on the standard template, INCLUDING
# their References formatting. So this is a search-and-extract, not a strict per-line match --
# find the ID token wherever it is, take the version if there is one, treat everything else as name.
COMPONENT_ID_RE = re.compile(r"\bTMFC\d+\b")
API_ID_RE = re.compile(r"\bTMF(?!C)\d+\b")  # (?!C) -- don't let TMFC020 match as if it were an API id
VERSION_RE = re.compile(r"\bv?(\d+\.\d+(?:\.\d+)?)\b")
unparsed_refs = []  # bullets under a recognized block with no ID found -- surfaced as a warning, not dropped silently


def clean_name(rest):
    name = re.sub(r"[-–—()]+", " ", rest)
    return re.sub(r"\s{2,}", " ", name).strip(" -–—:")


def parse_component_bullet(plain):
    m = COMPONENT_ID_RE.search(plain)
    if not m:
        return None
    rest = plain[: m.start()] + plain[m.end() :]
    vm = VERSION_RE.search(rest)
    version = vm.group(1) if vm else None
    if vm:
        rest = rest[: vm.start()] + rest[vm.end() :]
    entry = {"id": m.group(0), "name": clean_name(rest)}
    if version:
        entry["spec_version"] = version
    return entry


def parse_api_bullet(plain):
    m = API_ID_RE.search(plain)
    if not m:
        return None
    rest = plain[: m.start()] + plain[m.end() :]
    vm = VERSION_RE.search(rest)
    if vm:
        api_version = "v" + vm.group(1)
        rest = rest[: vm.start()] + rest[vm.end() :]
    else:
        api_version = None
    entry = {"id": m.group(0), "name": clean_name(rest)}
    if api_version:
        entry["api_version"] = api_version
    return entry

for kind, el in flatten(body):
    if kind == "p":
        text, style_id, is_list = paragraph_text_and_style(el)
        level = HEADING_LEVEL.get(style_id)

        if not seen_first_heading:
            # still in the title block -- never emitted, but scan it for name/id/version.
            # Some templates start with an empty Heading1-styled paragraph (a spacer/anchor,
            # not a real chapter heading) -- only a heading with actual text ends the title block.
            if level is not None and strip_md(text):
                seen_first_heading = True
            else:
                title_text = strip_md(text)
                m = re.match(r"^Use Case:\s*(.+)$", title_text)
                if not m:
                    # Introductory Guide / Guidebook / newer Use-Case cover pages don't use the
                    # "Use Case: <name>" phrasing at all -- some put "<ID> <name>" on one line
                    # (TMFS004), others put the name on its own line with the bare ID appearing
                    # separately, further down (TMFS006/010/019A/019B). Cover-page boilerplate like
                    # "TM Forum Introductory Guide" sits between them, so track the most recent
                    # non-boilerplate line as a candidate and commit it once the bare-ID line is
                    # actually reached -- that's the one part of this shape that's unambiguous.
                    m = re.match(rf"^{re.escape(ID)}\s+(.+)$", title_text)
                if m and envelope["name"] is None:
                    envelope["name"] = m.group(1)
                elif envelope["name"] is None and title_text:
                    # TMFS019A/TMFS019B are this repo's own split of a single TM Forum identifier
                    # (spec/tasks.md 4.2) -- the DOCX itself only ever says the bare "TMFS019".
                    if title_text in (ID, re.sub(r"[A-Z]$", "", ID)):
                        if title_candidate:
                            envelope["name"] = title_candidate
                    elif not title_text.startswith("TM Forum "):
                        title_candidate = title_text
                continue

        if level is not None:
            name = strip_md(text).lower()
            if drop_until_level is not None and level <= drop_until_level:
                drop_until_level = None
                in_references = False
                ref_block = None
            if drop_until_level is None and name in DROP_HEADINGS:
                drop_until_level = level
                in_references = name == "references"
                ref_block = None
                continue
            if drop_until_level is not None:
                continue  # nested heading inside an already-dropped subtree
            lines.append(paragraph_markdown(text, style_id, is_list))
            continue

        if drop_until_level is not None:
            if in_references:
                raw_plain = strip_md(text)
                # A single list-numbered paragraph can itself bundle several entries via manual
                # line breaks (w:br) instead of separate list items (confirmed in TMFS028) --
                # python-docx flattens w:br to "\n" in run.text, so without this split the first
                # ID's "rest" swallows every following line (including other IDs) into one `name`,
                # producing a multi-line YAML scalar that breaks the frontmatter parser outright.
                for plain in (raw_plain.split("\n") if "\n" in raw_plain else [raw_plain]):
                    plain = plain.strip()
                    low = plain.lower()
                    # Primary: search every (sub-)line in the References subtree for a component/API
                    # ID, regardless of list status or which (if any) sub-header it falls under.
                    # Confirmed necessary against the full 32-use-case set (Phase 4), not just the
                    # six-doc pilot -- header phrasing varies far more than "ODA components:" /
                    # "Open APIs:": some templates use "TMF references" or "The following documents
                    # are referenced..." as the header, and some format "ODA Components:" itself as a
                    # list bullet rather than plain prose (TMFS008) -- either way, the old
                    # header-keyword-then-bullet state machine silently dropped every real TMFC/TMF
                    # entry underneath. Matching the ID pattern directly, independent of header
                    # recognition, is what the "search-based, not strict-format" principle above
                    # (COMPONENT_ID_RE) already called for.
                    if COMPONENT_ID_RE.search(plain):
                        envelope["links"]["components"].append(parse_component_bullet(plain))
                    elif API_ID_RE.search(plain):
                        envelope["links"]["apis"].append(parse_api_bullet(plain))
                    elif not is_list:
                        # a header/narrative line -- keyword-classify by substring (phrasing varies)
                        # so a *list* bullet with no ID of its own can still be surfaced as
                        # "unparsed" below rather than silently dropped; blank lines don't reset it.
                        if plain:
                            if low.startswith("sid:") or low == "sid":
                                ref_block = "sid"
                            elif "component" in low:
                                ref_block = "components"
                            elif "open api" in low or "apis" in low:
                                ref_block = "apis"
                            else:
                                ref_block = None  # unrecognized header -- don't misattribute bullets
                    elif ref_block == "sid":
                        extension["sid_references"].append(plain)
                    elif ref_block in ("components", "apis") and plain:
                        unparsed_refs.append(("component" if ref_block == "components" else "api", plain))
            continue  # dropped, whether or not it was a References line

        md = paragraph_markdown(text, style_id, is_list)
        if md == "" and "![" not in text:
            lines.append("")  # blank paragraph -> blank line
        else:
            lines.append(md)

    elif kind == "tbl":
        md, cells_flat = table_markdown(el)
        if not seen_first_heading:
            for c in cells_flat:
                m = re.match(r"^Version\s+(\S+)$", strip_md(c))
                if m and envelope["version"] is None:
                    envelope["version"] = m.group(1)
            continue
        if drop_until_level is not None:
            continue
        lines.append(md)
        lines.append("")

body_md = "\n\n".join(lines)
body_md = re.sub(r"\n{3,}", "\n\n", body_md)
prev = None
while prev != body_md:
    prev = body_md
    body_md = body_md.replace("****", "")  # merge adjacent same-style runs Word split apart


frontmatter_data = {**envelope, **extension}  # envelope fields first, then extensions -- spec.md 5.1's documented order
yaml_lite.write(OUT, frontmatter_data, body_md)

print(
    f"Wrote {OUT} ({len(body_md)} chars body, {img_counter} images, "
    f"{len(envelope['links']['components'])} components, {len(envelope['links']['apis'])} apis linked)"
)
if envelope["name"] is None or envelope["version"] is None:
    print("WARNING: could not parse name/version from the title block -- check the DOCX structure")
if unparsed_refs:
    print(f"WARNING: {len(unparsed_refs)} References bullet(s) had no recognizable ID -- not linked, check by hand:")
    for kind, plain in unparsed_refs:
        print(f"    [{kind}] {plain!r}")
if not envelope["links"]["components"] and not envelope["links"]["apis"] and not unparsed_refs:
    print("WARNING: no components or APIs linked at all -- this document's References section may use a format this script doesn't recognize yet; check by hand")
